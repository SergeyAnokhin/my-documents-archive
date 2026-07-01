"""Pins native .docx text extraction — see docs/code-map.md (services/docx_extract.py).

extract_docx_text() concatenates paragraph and table text in document order,
mirroring the "\n\n"-per-block joining convention services/ocr.py uses, so
downstream consumers (search, analysis, embeddings) see a familiar text shape.
"""
from docx import Document as DocxDocument

from app.services.docx_extract import extract_docx_text


def test_extract_docx_text_concatenates_paragraphs(tmp_path):
    # Doc: services/docx_extract.py
    # Rule: non-empty paragraphs are joined with a blank line between them;
    #       empty/whitespace-only paragraphs are dropped.
    path = tmp_path / "a.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("   ")
    doc.add_paragraph("Second paragraph.")
    doc.save(path)

    text = extract_docx_text(str(path))
    assert text == "First paragraph.\n\nSecond paragraph."


def test_extract_docx_text_includes_table_cells(tmp_path):
    # Doc: services/docx_extract.py
    # Rule: table rows are extracted as pipe-joined cell text, appended after
    #       paragraph text, in document order.
    path = tmp_path / "b.docx"
    doc = DocxDocument()
    doc.add_paragraph("Intro.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    doc.save(path)

    text = extract_docx_text(str(path))
    assert text == "Intro.\n\nName | Value"


def test_extract_docx_text_empty_document_returns_empty_string(tmp_path):
    # Doc: services/docx_extract.py
    # Rule: a docx with no text content returns "" (falsy), not None or an
    #       error — callers rely on this for the "no text" analysis-skip check.
    path = tmp_path / "empty.docx"
    DocxDocument().save(path)
    assert extract_docx_text(str(path)) == ""
