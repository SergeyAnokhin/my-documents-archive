"""
Native .docx text extraction — no OCR involved (text is already digital).

Used by the indexer as an OCR-equivalent step for Word documents: paragraphs
and table cells are concatenated in document order, "\n\n"-joined per block
(matches services/ocr.py's block-join convention) so downstream consumers
(fulltext search, embeddings, AI analysis) see a familiar text shape.

Out of scope: headers/footers, footnotes, textboxes, comments — python-docx's
.paragraphs/.tables only walk the main document body, which is enough for
search/analysis purposes.
"""

import logging
from pathlib import Path

log = logging.getLogger(__name__)


def extract_docx_text(filepath: str) -> str:
    """Return concatenated paragraph + table text from a .docx file.

    Raises on unreadable/corrupt/encrypted files — caller (indexer) is
    responsible for catching and marking ocr_status=error, same contract as
    ocr.py::extract_text().
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(Path(filepath))
    blocks: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            blocks.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n\n".join(blocks)
